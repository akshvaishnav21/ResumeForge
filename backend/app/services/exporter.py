import re
import os
from io import BytesIO


def to_markdown(content: str) -> bytes:
    return content.encode("utf-8")


def to_txt(content: str) -> bytes:
    text = re.sub(r'#{1,6}\s+', '', content)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text.encode("utf-8")


def _strip_bold(text: str) -> str:
    return re.sub(r'\*\*(.*?)\*\*', r'\1', text)


def _find_font():
    """Find a Unicode TTF font on the system."""
    candidates = [
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def _find_font_bold():
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "C:/Windows/Fonts/segoeuib.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def to_pdf(content: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF('P', 'mm', 'Letter')
    pdf.set_auto_page_break(auto=True, margin=8)

    # Register Unicode font
    font_path = _find_font()
    font_bold_path = _find_font_bold()
    if font_path:
        pdf.add_font('Resume', '', font_path, uni=True)
        if font_bold_path:
            pdf.add_font('Resume', 'B', font_bold_path, uni=True)
        else:
            pdf.add_font('Resume', 'B', font_path, uni=True)
        FONT = 'Resume'
    else:
        FONT = 'Helvetica'

    pdf.add_page()
    pdf.set_left_margin(12)
    pdf.set_right_margin(12)
    pdf.set_y(8)

    LH = 3.2  # line height
    FS = 8    # font size

    def write_rich(text, font_size=FS, line_h=LH):
        """Write text with **bold** support and [link](url) hyperlinks."""
        # Split into bold, link, and plain segments
        tokens = re.split(r'(\*\*.*?\*\*|\[[^\]]+\]\([^\)]+\))', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                pdf.set_font(FONT, 'B', font_size)
                pdf.write(line_h, token[2:-2])
            elif token.startswith('['):
                m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', token)
                if m:
                    label, url = m.group(1), m.group(2)
                    pdf.set_font(FONT, '', font_size)
                    pdf.set_text_color(30, 80, 180)
                    pdf.write(line_h, label, url)
                    pdf.set_text_color(0, 0, 0)
                else:
                    pdf.set_font(FONT, '', font_size)
                    pdf.write(line_h, token)
            else:
                pdf.set_font(FONT, '', font_size)
                pdf.write(line_h, token)

    def write_contact_line(text, line_h=LH):
        """Write contact line with auto-detected URLs and emails as hyperlinks."""
        # Split by | separator, handle each segment
        segments = text.split('|')
        for i, seg in enumerate(segments):
            if i > 0:
                pdf.set_font(FONT, '', FS)
                pdf.write(line_h, ' | ')
            seg = seg.strip()
            # Check for markdown link
            md_link = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', seg)
            if md_link:
                label, url = md_link.group(1), md_link.group(2)
                pdf.set_font(FONT, '', FS)
                pdf.set_text_color(30, 80, 180)
                pdf.write(line_h, label, url)
                pdf.set_text_color(0, 0, 0)
            elif re.match(r'https?://', seg):
                pdf.set_font(FONT, '', FS)
                pdf.set_text_color(30, 80, 180)
                pdf.write(line_h, seg, seg)
                pdf.set_text_color(0, 0, 0)
            elif '@' in seg and '.' in seg:
                pdf.set_font(FONT, '', FS)
                pdf.set_text_color(30, 80, 180)
                pdf.write(line_h, seg, f'mailto:{seg}')
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.set_font(FONT, '', FS)
                pdf.write(line_h, seg)

    lines = content.split('\n')
    for line in lines:
        line = line.rstrip()

        if not line:
            pdf.ln(1)
            continue

        # H1 — Name
        if line.startswith('# '):
            text = _strip_bold(line[2:])
            pdf.set_font(FONT, 'B', 14)
            pdf.cell(0, 6, text, ln=True, align='C')
            continue

        # H2 — Section
        if line.startswith('## '):
            text = _strip_bold(line[3:])
            pdf.ln(2.0)
            pdf.set_font(FONT, 'B', 9.5)
            pdf.cell(0, 4, text.upper(), ln=True)
            y = pdf.get_y()
            pdf.set_draw_color(150, 150, 150)
            pdf.set_line_width(0.2)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(1)
            continue

        # H3 — Job / Company (split layout: left bold, right-aligned dates)
        if line.startswith('### '):
            text = _strip_bold(line[4:])
            pdf.ln(1)
            parts = [p.strip() for p in text.split('|')]
            if len(parts) >= 3:
                company, role = parts[0], parts[1]
                right_text = ' | '.join(parts[2:])
                usable_w = pdf.w - pdf.l_margin - pdf.r_margin

                pdf.set_font(FONT, '', 8)
                right_w = pdf.get_string_width(right_text) + 1

                pdf.set_font(FONT, 'B', 8.5)
                left_text = f"{company} \u2014 {role}"
                pdf.cell(usable_w - right_w, LH, left_text, ln=False)

                pdf.set_font(FONT, '', 8)
                pdf.cell(right_w, LH, right_text, ln=True, align='R')
            else:
                pdf.set_font(FONT, 'B', 8.5)
                pdf.cell(0, LH, text, ln=True)
            continue

        # Bullet
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:]
            pdf.set_x(pdf.l_margin)
            pdf.set_font(FONT, '', FS)
            pdf.cell(3, LH, '\u2022' if FONT != 'Helvetica' else '-', ln=False)
            write_rich(bullet_text)
            pdf.ln(LH + 0.3)
            continue

        # Bold area label like **Area Name:** or **Area Name**
        if re.match(r'^\*\*.*\*\*:?\s*$', line):
            text = _strip_bold(line).rstrip(':').strip()
            pdf.ln(1.0)
            pdf.set_font(FONT, 'B', FS)
            pdf.cell(0, LH, text + ':', ln=True)
            continue

        # Contact line or plain text
        pdf.set_font(FONT, '', FS)
        plain = _strip_bold(line)
        if '|' in plain or '@' in plain:
            # Center the contact line with clickable links
            # Measure total width using display labels (not raw URLs)
            segments = plain.split('|')
            pdf.set_font(FONT, '', FS)
            total_w = 0
            for i, seg in enumerate(segments):
                display = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', seg.strip())
                total_w += pdf.get_string_width(display)
                if i > 0:
                    total_w += pdf.get_string_width(' | ')
            start_x = max(pdf.l_margin, (pdf.w - total_w) / 2)
            pdf.set_x(start_x)
            write_contact_line(plain)
            pdf.ln(LH)
        else:
            write_rich(line)
            pdf.ln(LH)

    buf = BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def to_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    style = doc.styles['Normal']
    style.font.size = Pt(8.5)
    style.font.name = 'Calibri'
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0.5)
    style.paragraph_format.line_spacing = 1.0

    lines = content.split('\n')
    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_strip_bold(line[2:]))
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(1)

        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(_strip_bold(line[3:]).upper())
            run.bold = True
            run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '888888'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(_strip_bold(line[4:]))
            run.bold = True
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0.5)

        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0.3)
            p.paragraph_format.line_spacing = 1.0
            _add_rich_runs(p, line[2:], size=8.5)

        elif re.match(r'^\*\*.*\*\*:?\s*$', line):
            p = doc.add_paragraph()
            text = _strip_bold(line).rstrip(':').strip()
            run = p.add_run(text + ':')
            run.bold = True
            run.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(0.3)

        else:
            p = doc.add_paragraph()
            if '|' in line or '@' in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_runs(p, line, size=8.5)
            p.paragraph_format.space_after = Pt(0.5)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_hyperlink(paragraph, text, url, size=8.5):
    """Add a clickable hyperlink to a docx paragraph."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    import docx.oxml

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = docx.oxml.OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = docx.oxml.OxmlElement('w:r')
    rPr = docx.oxml.OxmlElement('w:rPr')
    c = docx.oxml.OxmlElement('w:color')
    c.set(qn('w:val'), '1E50B4')
    rPr.append(c)
    u = docx.oxml.OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = docx.oxml.OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_rich_runs(paragraph, text: str, size: float = 8.5):
    from docx.shared import Pt
    # Split into bold, links, and plain segments
    parts = re.split(r'(\*\*.*?\*\*|\[[^\]]+\]\([^\)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', part)
            if m:
                _add_hyperlink(paragraph, m.group(1), m.group(2), size)
            else:
                run = paragraph.add_run(part)
                run.font.size = Pt(size)
        else:
            # Auto-detect bare URLs and emails in contact lines
            sub_parts = re.split(r'(https?://\S+|\S+@\S+\.\S+)', part)
            for sp in sub_parts:
                if not sp:
                    continue
                if re.match(r'https?://', sp):
                    _add_hyperlink(paragraph, sp, sp, size)
                elif '@' in sp and '.' in sp and re.match(r'\S+@\S+\.\S+', sp):
                    _add_hyperlink(paragraph, sp, f'mailto:{sp}', size)
                else:
                    run = paragraph.add_run(sp)
                    run.font.size = Pt(size)
